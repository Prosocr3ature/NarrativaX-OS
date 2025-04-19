import os
import json
import requests
import zipfile
import random
import replicate
import threading
import queue
import base64
import time
from html import escape
from concurrent.futures import ThreadPoolExecutor, TimeoutError
from docx import Document
from docx.shared import Inches
from fpdf import FPDF
from tempfile import NamedTemporaryFile
from gtts import gTTS
from PIL import Image
from io import BytesIO
import streamlit as st
from streamlit.runtime.scriptrunner import add_script_run_ctx

# ========== INITIALIZATION ==========
st.set_page_config(
    page_title="NarrativaX", 
    page_icon="🪶", 
    layout="wide", 
    initial_sidebar_state="collapsed"
)

# ========== CONSTANTS ==========
LOGO_URL = "https://raw.githubusercontent.com/Prosocr3ature/NarrativaX/main/logo.png"
MAX_TOKENS = 1800
IMAGE_SIZE = (768, 1024)
PROGRESS_QUEUE = queue.Queue()
TIMEOUT = 300  # 5 minutes

SAFE_LOADING_MESSAGES = [
    "Sharpening quills...", "Mixing metaphorical ink...",
    "Convincing characters to behave...", "Battling clichés...",
    "Negotiating with plot holes...", "Summoning muses...",
    "Where we're going, we don't need chapters...",
    "Wordsmithing in progress...", "The pen is mightier... loading...",
    "Subverting expectations..."
]

TONE_MAP = {
    "Romantic": "sensual, romantic, literary",
    "Dark Romantic": "moody, passionate, emotional",
    "NSFW": "detailed erotic, emotional, mature",
    "Hardcore": "intense, vulgar, graphic, pornographic",
    "BDSM": "dominant, submissive, explicit, power-play",
    "Playful": "flirty, teasing, lighthearted",
    "Mystical": "dreamlike, surreal, poetic",
    "Gritty": "raw, realistic, street-style",
    "Slow Burn": "subtle, growing tension, emotional depth",
    "Wholesome": "uplifting, warm, feel-good",
    "Suspenseful": "tense, thrilling, page-turning",
    "Philosophical": "deep, reflective, thoughtful"
}

GENRES = [
    "Adventure", "Fantasy", "Dark Fantasy", "Romance", "Thriller",
    "Mystery", "Drama", "Sci-Fi", "Slice of Life", "Horror", "Crime",
    "LGBTQ+", "Action", "Psychological", "Historical Fiction",
    "Supernatural", "Steampunk", "Cyberpunk", "Post-Apocalyptic",
    "Surreal", "Noir", "Erotica", "NSFW", "Hardcore", "BDSM",
    "Futanari", "Incubus/Succubus", "Monster Romance",
    "Dubious Consent", "Voyeurism", "Yaoi", "Yuri", "Taboo Fantasy"
]

IMAGE_MODELS = {
    "Realistic Vision v5.1": "lucataco/realistic-vision-v5.1:2c8e954decbf70b7607a4414e5785ef9e4de4b8c51d50fb8b8b349160e0ef6bb",
    "Reliberate V3 (NSFW)": "asiryan/reliberate-v3:d70438fcb9bb7adb8d6e59cf236f754be0b77625e984b8595d1af02cdf034b29"
}

# ========== SESSION STATE ==========
for key in ['book', 'outline', 'cover', 'characters', 'gen_progress']:
    st.session_state.setdefault(key, None)
st.session_state.setdefault('image_cache', {})

# ========== CORE FUNCTIONS ==========
def pil_to_base64(image: Image.Image) -> str:
    buffered = BytesIO()
    image.save(buffered, format="PNG")
    return base64.b64encode(buffered.getvalue()).decode("utf-8")

def base64_to_pil(b64_str: str) -> Image.Image:
    return Image.open(BytesIO(base64.b64decode(b64_str)))

def call_openrouter(prompt: str, model: str) -> str:
    headers = {"Authorization": f"Bearer {os.getenv('OPENROUTER_API_KEY')}"}
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.95,
        "max_tokens": MAX_TOKENS
    }
    try:
        response = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=60
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        PROGRESS_QUEUE.put(("ERROR", f"API Error: {str(e)}", 0, ""))
        raise

def generate_image(prompt: str, model_key: str, id_key: str) -> Image.Image:
    try:
        if id_key in st.session_state.image_cache:
            cached = st.session_state.image_cache[id_key]
            return base64_to_pil(cached) if isinstance(cached, str) else cached

        output = replicate.run(
            IMAGE_MODELS[model_key],
            input={
                "prompt": f"{escape(prompt[:250])} {random.choice(['intricate details', 'cinematic lighting', '8k resolution'])}",
                "negative_prompt": "text, watermark, deformed, blurry",
                "num_inference_steps": 35,
                "width": IMAGE_SIZE[0],
                "height": IMAGE_SIZE[1]
            }
        )

        if output and isinstance(output, list):
            response = requests.get(output[0], timeout=30)
            response.raise_for_status()
            image = Image.open(BytesIO(response.content)).convert("RGB")
            st.session_state.image_cache[id_key] = pil_to_base64(image)
            return image
    except Exception as e:
        PROGRESS_QUEUE.put(("ERROR", f"Image Error: {str(e)}", 0, ""))
    return None

def background_generation_task():
    try:
        config = st.session_state.gen_progress
        total_steps = 4 + (config['chapters'] * 3)
        current_step = 0

        def heartbeat():
            while st.session_state.gen_progress:
                PROGRESS_QUEUE.put(("💓", "Processing...", current_step/total_steps, ""))
                time.sleep(10)

        heartbeat_thread = threading.Thread(target=heartbeat, daemon=True)
        add_script_run_ctx(heartbeat_thread)
        heartbeat_thread.start()

        # Phase 1: Concept Development
        PROGRESS_QUEUE.put(("🌌", "Developing core concept...", current_step/total_steps, ""))
        premise = call_openrouter(
            f"Develop a {config['genre']} story premise: {escape(config['prompt'])}",
            config['model']
        )
        current_step += 1

        # Phase 2: Outline Generation
        PROGRESS_QUEUE.put(("📜", "Crafting detailed outline...", current_step/total_steps, ""))
        outline_prompt = f"""Create detailed outline for {TONE_MAP[config['tone']]} {config['genre']} novel: {premise}
        Include chapter breakdowns, character arcs, and key plot points."""
        st.session_state.outline = call_openrouter(outline_prompt, config['model'])
        current_step += 1

        # Phase 3: Content Generation
        book = {}
        sections = ["Foreword"] + [f"Chapter {i+1}" for i in range(config['chapters'])] + ["Epilogue"]
        
        for sec in sections:
            PROGRESS_QUEUE.put(("📖", f"Writing {sec}...", current_step/total_steps, ""))
            content = call_openrouter(
                f"Write immersive '{sec}' content for {config['genre']} novel: {st.session_state.outline}",
                config['model']
            )
            book[sec] = content
            current_step += 1

            PROGRESS_QUEUE.put(("🎨", f"Generating {sec} image...", current_step/total_steps, ""))
            generate_image(f"{escape(content[:200])} {TONE_MAP[config['tone']]} style", config['img_model'], sec)
            current_step += 1

        # Phase 4: Final Assets
        PROGRESS_QUEUE.put(("🖼️", "Creating cover art...", current_step/total_steps, ""))
        st.session_state.cover = generate_image(
            f"Cinematic cover art for {config['genre']} novel: {premise}",
            config['img_model'],
            "cover"
        )
        current_step += 1

        PROGRESS_QUEUE.put(("👥", "Developing characters...", current_step/total_steps, ""))
        st.session_state.characters = json.loads(call_openrouter(
            f"""Generate characters for {config['genre']} novel in JSON format:
            {st.session_state.outline}
            Format: [{{"name":"","role":"","personality":"","appearance":""}}]""",
            config['model']
        ))
        current_step += 1

        st.session_state.book = book
        PROGRESS_QUEUE.put(("COMPLETE", "📖 Book generation complete!", 1.0, ""))

    except Exception as e:
        PROGRESS_QUEUE.put(("ERROR", f"Generation failed: {str(e)}", 0, ""))
        st.session_state.gen_progress = None

def background_generation_wrapper():
    try:
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(background_generation_task)
            future.result(timeout=TIMEOUT)
    except TimeoutError:
        PROGRESS_QUEUE.put(("ERROR", "Generation timed out after 5 minutes", 0, ""))
    finally:
        st.session_state.gen_progress = None

# ========== UI COMPONENTS ==========
def dramatic_logo():
    safe_message = escape(random.choice(SAFE_LOADING_MESSAGES))
    st.markdown(f"""
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@500&family=Cinzel:wght@500&display=swap');
        @keyframes float {{
            0% {{ transform: translate(-50%, -55%) rotate(-5deg); }}
            50% {{ transform: translate(-50%, -60%) rotate(5deg); }}
            100% {{ transform: translate(-50%, -55%) rotate(-5deg); }}
        }}
        .logo-overlay {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            height: 100%;
            background: rgba(0,0,0,0.9);
            z-index: 9998;
        }}
        .logo-container {{
            position: fixed;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            text-align: center;
            z-index: 9999;
        }}
        .logo-img {{
            width: min(80vw, 600px);
            animation: float 3.5s ease-in-out infinite;
            filter: drop-shadow(0 0 20px #ff69b480);
        }}
        .loading-message {{
            font-family: 'Playfair Display', serif;
            font-size: clamp(2rem, 5vw, 3rem);
            margin-top: 2rem;
            color: #ff69b4;
            animation: pulse 1.5s infinite;
        }}
        @keyframes pulse {{
            0% {{ opacity: 0.8; text-shadow: 0 0 10px #ff69b4; }}
            50% {{ opacity: 1; text-shadow: 0 0 20px #ff69b4; }}
            100% {{ opacity: 0.8; text-shadow: 0 0 10px #ff69b4; }}
        }}
        .progress-bar {{
            height: 15px !important;
            border-radius: 10px;
            background: #ffffff20;
        }}
        .progress-bar::-webkit-progress-value {{
            background: linear-gradient(90deg, #ff69b4, #ff1493);
            border-radius: 10px;
        }}
    </style>
    <div class="logo-overlay"></div>
    <div class="logo-container">
        <img class="logo-img" src="{LOGO_URL}">
        <div class="loading-message">{safe_message}</div>
    </div>
    """, unsafe_allow_html=True)

def progress_animation():
    try:
        if not PROGRESS_QUEUE.empty():
            status = PROGRESS_QUEUE.get()
            
            with st.empty() as container:
                while True:
                    if status[0] == "COMPLETE":
                        st.balloons()
                        st.session_state.gen_progress = None
                        break
                    elif status[0] == "ERROR":
                        st.error(f"🚨 {escape(str(status[1]))[:200]}...")
                        st.session_state.gen_progress = None
                        break
                    else:
                        emoji, message, progress, preview = status
                        safe_preview = escape(str(preview))[:150] + "..." if preview else ""
                        container.markdown(f"""
                        <div style="text-align: center; padding: 2rem">
                            <div style="font-size: 3rem; animation: pulse 1.5s infinite">{emoji}</div>
                            <h3 style="margin: 1rem 0">{escape(message)}</h3>
                            <progress class="progress-bar" value="{progress}" max="1"></progress>
                            {f'<div style="background: rgba(255,255,255,0.1); border-radius: 10px; padding: 1rem; margin: 1rem 0">{safe_preview}</div>' if preview else ''}
                        </div>
                        """, unsafe_allow_html=True)
                    
                    try:
                        status = PROGRESS_QUEUE.get(timeout=0.1)
                    except queue.Empty:
                        break
    except Exception as e:
        st.error(f"Animation Error: {escape(str(e))[:200]}...")
        st.session_state.gen_progress = None

def main_interface():
    try:
        if st.session_state.get('gen_progress'):
            dramatic_logo()
            progress_animation()
            
            time.sleep(0.1)
            st.experimental_rerun()
        else:
            st.markdown(f'<img src="{escape(LOGO_URL)}" width="300" style="float:right; margin:-50px -20px 0 0">', 
                      unsafe_allow_html=True)
            st.title("NarrativaX — Immersive AI Book Creator")
            
            with st.container():
                prompt = st.text_area("🖋️ Your Story Concept", height=120,
                                    placeholder="A forbidden romance between...")
                col1, col2, col3 = st.columns(3)
                genre = col1.selectbox("📖 Genre", GENRES)
                tone = col2.selectbox("🎨 Tone", list(TONE_MAP))
                chapters = col3.slider("📚 Chapters", 4, 30, 10)
                model = col1.selectbox("🤖 AI Model", ["nothingiisreal/mn-celeste-12b", "gryphe/mythomax-l2-13b"])
                img_model = col2.selectbox("🖼️ Image Model", list(IMAGE_MODELS))

                if st.button("🚀 Create Book"):
                    st.session_state.image_cache.clear()
                    st.session_state.cover = None
                    st.session_state.book = None
                    st.session_state.outline = None
                    st.session_state.characters = None
                    
                    st.session_state.gen_progress = {
                        "prompt": prompt, "genre": genre, "tone": tone,
                        "chapters": chapters, "model": model, "img_model": img_model
                    }
                    gen_thread = threading.Thread(target=background_generation_wrapper, daemon=True)
                    add_script_run_ctx(gen_thread)
                    gen_thread.start()
                    st.rerun()
                    
    except Exception as e:
        st.error(f"Application Error: {escape(str(e))[:200]}...")
        st.session_state.gen_progress = None
        st.stop()

def render_sidebar():
    try:
        with st.sidebar:
            st.markdown(f'<img src="{escape(LOGO_URL)}" width="200" style="margin-bottom:20px">', unsafe_allow_html=True)
            
            if st.button("💾 Save Project"):
                try:
                    save_data = {
                        'book': st.session_state.book,
                        'outline': st.session_state.outline,
                        'characters': st.session_state.characters,
                        'image_cache': st.session_state.image_cache,
                        'cover': pil_to_base64(st.session_state.cover) if st.session_state.cover else None
                    }
                    with open("session.narrx", "w") as f:
                        json.dump(save_data, f)
                    st.success("Project saved!")
                except Exception as e:
                    st.error(f"Save failed: {escape(str(e))[:200]}...")

            if st.button("📂 Load Project"):
                try:
                    with open("session.narrx", "r") as f:
                        data = json.load(f)
                    
                    st.session_state.book = data.get('book')
                    st.session_state.outline = data.get('outline')
                    st.session_state.characters = data.get('characters')
                    st.session_state.image_cache = {
                        k: base64_to_pil(v) if isinstance(v, str) else v 
                        for k, v in data.get('image_cache', {}).items()
                    }
                    if data.get('cover'):
                        st.session_state.cover = base64_to_pil(data['cover'])
                    st.success("Project loaded!")
                except Exception as e:
                    st.error(f"Load failed: {escape(str(e))[:200]}...")

            if st.session_state.book and st.button("📦 Export Book"):
                with st.spinner("Packaging your masterpiece..."):
                    try:
                        with NamedTemporaryFile(delete=False, suffix=".zip") as tmp:
                            with zipfile.ZipFile(tmp.name, 'w') as zipf:
                                # DOCX
                                doc = Document()
                                for sec, content in st.session_state.book.items():
                                    doc.add_heading(sec, level=1)
                                    doc.add_paragraph(content)
                                    if sec in st.session_state.image_cache:
                                        img = st.session_state.image_cache[sec]
                                        if isinstance(img, str):
                                            img = base64_to_pil(img)
                                        img_io = BytesIO()
                                        img.save(img_io, format='PNG')
                                        doc.add_picture(img_io, width=Inches(5))
                                doc.save("book.docx")
                                zipf.write("book.docx")
                                os.remove("book.docx")

                                # PDF
                                pdf = FPDF()
                                pdf.set_auto_page_break(auto=True, margin=15)
                                if st.session_state.cover:
                                    cover_path = "cover.png"
                                    st.session_state.cover.save(cover_path)
                                    pdf.image(cover_path, x=0, y=0, w=pdf.w, h=pdf.h)
                                    pdf.add_page()
                                pdf.set_font("Arial", size=12)
                                for sec, content in st.session_state.book.items():
                                    pdf.set_font("Arial", 'B', 16)
                                    pdf.cell(0, 10, sec, ln=True)
                                    pdf.set_font("Arial", size=12)
                                    pdf.multi_cell(0, 10, content)
                                pdf.output("book.pdf")
                                zipf.write("book.pdf")
                                os.remove("book.pdf")

                                # Audio
                                for i, (sec, content) in enumerate(st.session_state.book.items()):
                                    with NamedTemporaryFile(delete=False, suffix=".mp3") as audio_tmp:
                                        tts = gTTS(text=content, lang='en')
                                        tts.save(audio_tmp.name)
                                        zipf.write(audio_tmp.name, f"chapter_{i+1}.mp3")
                                        os.remove(audio_tmp.name)

                            with open(tmp.name, "rb") as f:
                                st.download_button("⬇️ Download ZIP", f.read(), "narrativax_book.zip")
                            os.remove(tmp.name)
                    except Exception as e:
                        st.error(f"Export failed: {escape(str(e))[:200]}...")
    except Exception as e:
        st.error(f"Sidebar Error: {escape(str(e))[:200]}...")

def display_content():
    try:
        if st.session_state.book:
            st.header("Your Generated Book")
            
            with st.expander("📔 Book Cover", expanded=True):
                if st.session_state.cover:
                    st.image(st.session_state.cover, use_container_width=True)
                else:
                    st.warning("No cover generated yet")
            
            with st.expander("📝 Full Outline"):
                st.markdown(f"```\n{escape(st.session_state.outline)}\n```")
            
            with st.expander("👥 Character Bios"):
                for char in st.session_state.characters:
                    with st.container():
                        cols = st.columns([1, 3])
                        cols[0].subheader(escape(char.get('name', 'Unnamed')))
                        cols[1].write(f"""
                        **Role:** {escape(char.get('role', 'Unknown'))}  
                        **Personality:** {escape(char.get('personality', 'N/A'))}  
                        **Appearance:** {escape(char.get('appearance', 'Not specified'))}
                        """)

            for section, content in st.session_state.book.items():
                with st.expander(f"📜 {escape(section)}"):
                    col1, col2 = st.columns([3, 2])
                    with col1:
                        st.write(escape(content))
                        with NamedTemporaryFile(suffix=".mp3") as tf:
                            tts = gTTS(text=content, lang='en')
                            tts.save(tf.name)
                            st.audio(tf.name, format="audio/mp3")
                    with col2:
                        if section in st.session_state.image_cache:
                            img = st.session_state.image_cache[section]
                            if isinstance(img, str):
                                img = base64_to_pil(img)
                            st.image(img, use_container_width=True)
                        else:
                            st.warning("No image for this section")
    except Exception as e:
        st.error(f"Content Error: {escape(str(e))[:200]}...")

# ========== MAIN EXECUTION ==========
if __name__ == "__main__":
    main_interface()
    render_sidebar()
    display_content()
    st.markdown("""
    <style>
        @media (max-width: 768px) {
            .main-column { padding: 0 5px !important; }
            .stTextArea textarea { font-size: 16px !important; }
        }
        .stProgress > div > div > div { 
            background-color: #ff69b4 !important;
        }
    </style>
    """, unsafe_allow_html=True)
