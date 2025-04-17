import os, time, json, requests
import streamlit as st
from docx import Document
from fpdf import FPDF
from tempfile import NamedTemporaryFile
from gtts import gTTS
import replicate
from docx.shared import Inches

# --- CONFIG ---
st.set_page_config(page_title="NarrativaX", page_icon="🪶", layout="wide")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
REPLICATE_API_TOKEN = os.getenv("REPLICATE_API_TOKEN")
replicate_client = replicate.Client(api_token=REPLICATE_API_TOKEN)

TONE_MAP = {
    "Romantic": "sensual, romantic, literary",
    "Dark Romantic": "moody, passionate, emotional",
    "NSFW": "detailed erotic, emotional, mature",
    "Hardcore": "intense, vulgar, graphic, pornographic",
    "BDSM": "dominant, submissive, explicit, power-play",
    "Playful": "flirty, teasing, lighthearted",
    "Mystical": "dreamlike, surreal, poetic",
    "Gritty": "raw, realistic, street-style",
    "Slow Burn": "subtle, growing tension, emotional depth"
}
GENRES_NORMAL = [
    "Adventure", "Fantasy", "Dark Fantasy", "Romance", "Thriller", "Historical Fiction", "Mystery",
    "Drama", "Sci-Fi", "Slice of Life", "Teen Fiction", "Horror", "Cyberpunk", "Psychological", "Crime", "LGBTQ+", "Action", "Paranormal"
]
GENRES_ADULT = [
    "Erotica", "NSFW", "Hardcore", "BDSM", "Yaoi", "Yuri", "Futanari", "Harem", "Incubus/Succubus", "Kinky Comedy", "Taboo Fantasy"
]
GENRES = GENRES_NORMAL + GENRES_ADULT
VOICES = {"Rachel": "default", "Bella": "default", "Antoni": "default", "Elli": "default", "Josh": "default"}
MODELS = [
    "nothingiisreal/mn-celeste-12b",
    "openchat/openchat-3.5-0106",
    "gryphe/mythomax-l2-13b",
    "nousresearch/nous-capybara-7b",
    "cognitivecomputations/dolphin-mixtral"
]
IMAGE_MODELS = {
    "Realistic Vision v5.1": "lucataco/realistic-vision-v5.1:2c8e954decbf70b7607a4414e5785ef9e4de4b8c51d50fb8b8b349160e0ef6bb",
    "Reliberate V3 (NSFW)": "asiryan/reliberate-v3:d70438fcb9bb7adb8d6e59cf236f754be0b77625e984b8595d1af02cdf034b29"
}
SAFE_IMAGE_MODELS = {k: v for k, v in IMAGE_MODELS.items() if "NSFW" not in k}

# --- INIT STATE ---
def init_state():
    defaults = {
        "book": {}, "outline": "", "characters": [], "prompt": "",
        "genre": "", "tone": "", "adult_confirmed": False,
        "chapter_order": [], "image_cache": {}, "audio_cache": {},
        "img_model": "", "book_title": "", "custom_title": "", "tagline": "", "cover_image": None
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v
init_state()

# --- SAFETY ---
def is_adult_mode():
    return st.session_state.genre in GENRES_ADULT or st.session_state.tone in ["NSFW", "Hardcore", "BDSM"]

def require_adult_confirmation():
    if not st.session_state.adult_confirmed:
        with st.expander("⚠️ Adult Content", expanded=True):
            st.error("This content may contain mature or explicit material.")
            if st.button("I confirm I'm 18+"):
                st.session_state.adult_confirmed = True
            st.stop()

def check_adult_flag():
    if is_adult_mode() and not st.session_state.adult_confirmed:
        require_adult_confirmation()

# --- API FUNKTIONER ---
def call_openrouter(prompt, model, max_tokens=1800):
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://narrativax.app",
        "X-Title": "NarrativaX"
    }
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.95, "max_tokens": max_tokens
    }
    r = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"].strip()

def generate_outline(prompt, genre, tone, chapters, model):
    return call_openrouter(
        f"You are a ghostwriter. Create an outline for a {tone} {genre} novel with {chapters} chapters. "
        f"Include: Title, Foreword, Introduction, Chapter Titles, Final Words. Concept:\n{prompt}", model)

def generate_section(title, outline, model):
    return call_openrouter(f"Write section '{title}' in full based on this outline:\n{outline}", model)

def generate_characters(outline, genre, tone, model):
    prompt = f"""Create a list of characters for a {tone} {genre} novel based on the outline below. 
    For each, return JSON: name, role, personality, appearance.
    Output format: [{"name": "X", "role": "Y", "personality": "...", "appearance": "..."}]
    
    Outline: {outline}"""
    response = call_openrouter(prompt, model)
    try:
        return json.loads(response)
    except:
        st.warning("Could not parse characters. Showing raw format.")
        return [{"name": "Unnamed", "role": "Unknown", "personality": response, "appearance": ""}]

def generate_image(prompt, model_key, id_key):
    check_adult_flag()
    if id_key in st.session_state.image_cache:
        return st.session_state.image_cache[id_key]
    model = IMAGE_MODELS[model_key]
    args = {"prompt": prompt[:300], "num_inference_steps": 30, "guidance_scale": 7.5, "width": 768, "height": 1024}
    image_url = replicate_client.run(model, input=args)[0]
    st.session_state.image_cache[id_key] = image_url
    return image_url

def narrate(text, id_key):
    if id_key in st.session_state.audio_cache:
        return st.session_state.audio_cache[id_key]
    filename = f"{id_key}.mp3"
    gTTS(text.replace("\n", " ")).save(filename)
    st.session_state.audio_cache[id_key] = filename
    return filename

# --- SIDEBAR ---
with st.sidebar:
    st.image("https://i.imgur.com/vGV9N5k.png", width=180)
    st.markdown("**NarrativaX 5.0** — AI Book Studio")
    if st.button("Save Project"):
        json.dump(st.session_state.book, open("session.json", "w"))
        st.success("Project saved.")
    if st.button("Load Project"):
        try:
            st.session_state.book = json.load(open("session.json"))
            st.session_state.chapter_order = list(st.session_state.book.keys())
            st.success("Project loaded.")
        except Exception as e:
            st.error(f"Load failed: {e}")

    # Kapitelöversikt
    if st.session_state.chapter_order:
        st.markdown("### Chapters")
        for i, ch in enumerate(st.session_state.chapter_order):
            if st.button(f"Go to: {ch}", key=f"sidebar_ch_{i}"):
                st.session_state["jump_to_chapter"] = ch

# --- MAIN UI ---
st.title("NarrativaX — AI Book Studio")

# Visa omslag
if st.session_state.cover_image:
    st.image(st.session_state.cover_image, caption=f"**{st.session_state.book_title}**\n{st.session_state.tagline}", use_column_width=True)
    st.download_button("Download Cover", requests.get(st.session_state.cover_image).content, file_name="cover.jpg")

# --- SETTINGS ---
with st.expander("Book Settings", expanded=True):
    st.session_state.prompt = st.text_area("Book Idea / Prompt", height=150)
    genre_type = st.radio("Content Type", ["Normal", "Adult"], horizontal=True)
    genre_list = GENRES_ADULT if genre_type == "Adult" else GENRES_NORMAL
    st.session_state.genre = st.selectbox("Genre", genre_list)
    st.session_state.tone = st.selectbox("Tone", list(TONE_MAP))
    chapters = st.slider("Chapters", 4, 20, 10)
    model = st.selectbox("Model", MODELS)
    voice = st.selectbox("Voice", list(VOICES))
    st.session_state.img_model = st.selectbox("Image Model", list(IMAGE_MODELS if is_adult_mode() else SAFE_IMAGE_MODELS))
    st.session_state.custom_title = st.text_input("Custom Title (optional)", "")
    st.session_state.tagline = st.text_input("Tagline (optional)", "")

# --- GENERATE BOOK ---
if st.button("Create Full Book"):
    check_adult_flag()
    with st.spinner("Creating outline and characters..."):
        st.session_state.outline = generate_outline(st.session_state.prompt, st.session_state.genre, TONE_MAP[st.session_state.tone], chapters, model)
        st.session_state.characters = generate_characters(st.session_state.outline, st.session_state.genre, TONE_MAP[st.session_state.tone], model)

        # Bokomslag
        title_line = next((line for line in st.session_state.outline.splitlines() if "Title:" in line), None)
        raw_title = title_line.replace("Title:", "").strip() if title_line else "Untitled"
        st.session_state.book_title = st.session_state.custom_title or raw_title
        cover_prompt = f"{st.session_state.book_title}, {st.session_state.genre}, {st.session_state.tone}, book cover, centered, cinematic, ultra detailed"
        st.session_state.cover_image = generate_image(cover_prompt, st.session_state.img_model, "cover")

    with st.spinner("Writing full book..."):
        book = {}
        sections = ["Foreword", "Introduction"] + [f"Chapter {i+1}" for i in range(chapters)] + ["Final Words"]
        st.session_state.chapter_order = sections
        for section in sections:
            st.info(f"Writing {section}...")
            book[section] = generate_section(section, st.session_state.outline, model)
            img_prompt = f"Illustration for section '{section}': {book[section][:300]}"
            st.session_state.image_cache[section] = generate_image(img_prompt, st.session_state.img_model, section)
        st.session_state.book = book
        st.success("Done!")

# --- VISA BOK ---
if st.session_state.book:
    tab_titles = st.session_state.chapter_order + ["Characters"]
    active_index = next((i for i, t in enumerate(tab_titles) if t == st.session_state.get("jump_to_chapter")), 0)
    tabs = st.tabs(tab_titles)

    for i, title in enumerate(st.session_state.chapter_order):
        with tabs[i]:
            st.subheader(title)
            st.markdown(st.session_state.book[title])
            if title in st.session_state.image_cache:
                st.image(st.session_state.image_cache[title], caption=f"{title} Illustration")
            if st.button(f"Read Aloud: {title}", key=f"tts_{title}"):
                mp3 = narrate(st.session_state.book[title], title)
                st.audio(mp3)
    # --- KARAKTÄRER ---
    with tabs[-1]:
        st.subheader("Character Gallery")

        # Sök + filter
        search_query = st.text_input("Search by name, role or traits...", "").lower()
        all_roles = list({c['role'] for c in st.session_state.characters})
        role_filter = st.selectbox("Filter by Role", ["All"] + sorted(all_roles))

        filtered_chars = [
            c for c in st.session_state.characters
            if (search_query in c['name'].lower()
                or search_query in c['role'].lower()
                or search_query in c['personality'].lower()
                or search_query in c['appearance'].lower())
            and (role_filter == "All" or c['role'] == role_filter)
        ]

        edited = st.data_editor(
            filtered_chars,
            num_rows="dynamic",
            use_container_width=True,
            key="char_table"
        )
        st.session_state.characters = edited

        for i, char in enumerate(edited):
            with st.expander(f"{char['name']} — {char['role']}"):
                prompt = st.text_input("Portrait prompt", f"{char['appearance']}, portrait", key=f"imgprompt_{i}")
                if st.button(f"Generate Portrait", key=f"charportrait_{i}"):
                    img = generate_image(prompt, st.session_state.img_model, f"char_{i}")
                    st.image(img, caption=char['name'])
                    image_path = f"{char['name'].replace(' ', '_')}_portrait.jpg"
                    with open(image_path, "wb") as f:
                        f.write(requests.get(img).content)
                        st.download_button(f"Download {char['name']}", open(f.name, "rb"), file_name=image_path)

# --- EXPORT ---
st.header("Export")
pdf_mode = st.radio("PDF Export Mode", ["Include Chapter Images", "Text Only"], horizontal=True)
col1, col2, col3 = st.columns(3)

with col1:
    if st.button("Export DOCX"):
        doc = Document()
        if st.session_state.cover_image:
            doc.add_picture(requests.get(st.session_state.cover_image, stream=True).raw, width=Inches(5.5))
            doc.add_paragraph(st.session_state.book_title).bold = True
            doc.add_paragraph(st.session_state.tagline)
        for t, txt in st.session_state.book.items():
            doc.add_heading(t, level=1)
            doc.add_paragraph(txt)
        f = NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(f.name)
        st.download_button("Download DOCX", open(f.name, "rb"), file_name="NarrativaX_Book.docx")

with col2:
    if st.button("Export PDF"):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        if st.session_state.cover_image:
            image_url = st.session_state.cover_image
            image_path = "cover.jpg"
            with open(image_path, "wb") as f:
                f.write(requests.get(image_url).content)
            try:
                pdf.image(image_path, w=180)
                pdf.ln(10)
            except:
                pass

        for title in st.session_state.chapter_order:
            if pdf_mode == "Include Chapter Images" and title in st.session_state.image_cache:
                image_url = st.session_state.image_cache[title]
                image_path = f"{title.replace(' ', '_')}.jpg"
                with open(image_path, "wb") as f:
                    f.write(requests.get(image_url).content)
                try:
                    pdf.image(image_path, w=170)
                except:
                    pass

            pdf.set_font("Arial", style="B", size=14)
            pdf.cell(200, 10, title, ln=True)
            pdf.set_font("Arial", size=12)
            for line in st.session_state.book[title].splitlines():
                pdf.multi_cell(0, 10, line)

        f = NamedTemporaryFile(delete=False, suffix=".pdf")
        name = "NarrativaX_Book_Images.pdf" if pdf_mode == "Include Chapter Images" else "NarrativaX_Book_TextOnly.pdf"
        st.download_button("Download PDF", open(f.name, "rb"), file_name=name)

with col3:
    if st.button("Export JSON"):
        j = json.dumps(st.session_state.book)
        st.download_button("Download JSON", j, file_name="NarrativaX_Book.json")

# --- CLEAR JUMP ---
if "jump_to_chapter" in st.session_state:
    del st.session_state["jump_to_chapter"]

# --- FOOTER ---
st.markdown("---")
st.caption("© 2025 NarrativaX | AI storytelling without limits.")
