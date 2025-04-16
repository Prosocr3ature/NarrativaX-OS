# Filename: app.py

import os, time, json, requests
import streamlit as st
from docx import Document
from fpdf import FPDF
from tempfile import NamedTemporaryFile
from gtts import gTTS
import replicate
from streamlit_sortables import sort_items

# --- CONFIG ---
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
REPLICATE_API_TOKEN = os.getenv("REPLICATE_API_TOKEN")
replicate_client = replicate.Client(api_token=REPLICATE_API_TOKEN)

VOICES = {"Rachel": "default", "Bella": "default", "Antoni": "default", "Elli": "default", "Josh": "default"}
TONE_MAP = {
    "Romantic": "sensual, romantic, literary",
    "NSFW": "detailed erotic, emotional, mature",
    "Hardcore": "intense, vulgar, graphic, pornographic"
}
MODELS = [
    "nothingiisreal/mn-celeste-12b",
    "openchat/openchat-3.5-0106",
    "gryphe/mythomax-l2-13b",
    "nousresearch/nous-capybara-7b",
    "cognitivecomputations/dolphin-mixtral"
]
IMAGE_MODELS = {
    "Reliberate V3 (Erotica/NSFW)": "asiryan/reliberate-v3:d70438fcb9bb",
    "Realistic Vision 5.1": "lucataco/realistic-vision-v5.1:2c8e95"
}
GENRES = [
    "Erotica", "Dark Fantasy", "Sci-Fi", "Romance", "Thriller", "Adventure", "Historical Fiction",
    "Mystery", "Fantasy", "Drama", "Slice of Life", "Teen Fiction", "Horror", "Cyberpunk",
    "Psychological", "Crime", "LGBTQ+", "Action", "Paranormal"
]

# --- SESSION INIT ---
def initialize_state():
    defaults = {
        "last_saved": None,
        "feedback_history": [],
        "characters": [],
        "chapter_order": [],
        "book": {},
        "outline": "",
        "adult_confirmed": False
    }
    for key, val in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = val

initialize_state()

# --- SAFETY ---
def require_adult_verification():
    if not st.session_state.adult_confirmed:
        with st.sidebar:
            st.warning("**This content may be explicit. You must be 18+ to continue.**")
            if st.button("I Confirm I'm 18+ and Understand"):
                st.session_state.adult_confirmed = True
                st.rerun()
            st.stop()

def is_adult_genre(genre):
    return genre in ["Erotica", "NSFW", "Hardcore"]

# --- API CALLS ---
def call_openrouter(prompt, model, max_tokens=1800):
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://narrativax.app",
        "X-Title": "NarrativaX"
    }
    payload = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.95,
        "max_tokens": max_tokens
    }
    r = requests.post("https://openrouter.ai/api/v1/chat/completions", headers=headers, json=payload)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"].strip()

# --- GENERATORS ---
def generate_outline(prompt, genre, tone, chapters, model):
    return call_openrouter(
        f"You are a ghostwriter. Create a complete outline for a {tone} {genre} novel with {chapters} chapters. "
        f"Include: Title, Foreword, Introduction, {chapters} chapter titles, Final Words. Concept: {prompt}",
        model)

def generate_section(title, outline, model):
    return call_openrouter(
        f"Write only the section '{title}' in full detail based on the following novel outline. "
        f"Do not include other sections or chapters.\n\nOutline:\n{outline}\n\nSection: {title}",
        model)

def generate_full_book(outline, chapters, model):
    book = {}
    sections = ["Foreword", "Introduction"] + [f"Chapter {i+1}" for i in range(chapters)] + ["Final Words"]
    progress = st.progress(0)
    for idx, sec in enumerate(sections):
        book[sec] = generate_section(sec, outline, model)
        progress.progress((idx + 1) / len(sections))
    st.session_state.chapter_order = sections.copy()
    return book

def generate_characters(prompt, genre, tone, model, count=3):
    result = call_openrouter(
        f"Generate {count} unique characters for a {tone} {genre} story based on this: {prompt}. "
        f"Format: Name, Role, Appearance, Personality, Motivation, Secret.", model)
    return result.split("\n\n")

def generate_image(prompt, model_key="Realistic Vision 5.1"):
    if is_adult_genre(st.session_state.get("genre", "")) and not st.session_state.adult_confirmed:
        st.warning("You must confirm your age to access adult image models.")
        return None
    model = IMAGE_MODELS.get(model_key)
    input_args = {
        "prompt": prompt.strip()[:300],
        "num_inference_steps": 30,
        "guidance_scale": 7.5,
        "width": 768,
        "height": 1024
    }
    output = replicate_client.run(model, input=input_args)
    return output[0]

def generate_cover(prompt, model_key="Realistic Vision 5.1"):
    return generate_image(prompt + ", full book cover, illustration", model_key)

def narrate_story(text, voice_id=None):
    tts = gTTS(text.replace("\n", " ").strip())
    filename = f"narration_{voice_id or 'default'}.mp3"
    tts.save(filename)
    return filename

# --- EXPORT ---
def export_docx(data):
    doc = Document()
    for k, v in data.items():
        doc.add_heading(k, level=1)
        doc.add_paragraph(v)
    f = NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(f.name)
    return f.name

def export_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    for k, v in data.items():
        pdf.set_font("Arial", style="B", size=14)
        pdf.cell(200, 10, k, ln=True)
        pdf.set_font("Arial", size=12)
        for line in v.splitlines():
            pdf.multi_cell(0, 10, line)
    f = NamedTemporaryFile(delete=False, suffix=".pdf")
    pdf.output(f.name)
    return f.name
